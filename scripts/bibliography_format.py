# -*- coding: utf-8 -*-
"""
This script automates the formatting of a DOCX dissertation file.

It performs two main tasks:
1.  Formats all footnotes to use Times New Roman 10pt font with justified
    alignment, 1.0 line spacing, and a consistent tab stop for alignment.
2.  Sorts the bibliography section based on language (Cyrillic then Latin)
    and categorizes entries according to metadata from a Zotero library.
"""

# --- Standard Library Imports ---
import os
import re
import shutil
import zipfile
from collections import defaultdict

# --- Third-Party Imports ---
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt
from lxml import etree
from pyzotero import zotero
from transliterate import translit

# ==============================================================================
# --- USER SETTINGS ---
# ==============================================================================

# --- Zotero API Configuration ---
# You can find your User ID and generate an API Key here:
# https://www.zotero.org/settings/keys
API_KEY = 'YOUR_ZOTERO_API_KEY'
USER_ID = 'YOUR_ZOTERO_USER_ID'

# --- Zotero Library and Collection ---
# Set to 'user' for personal libraries or 'group' for group libraries.
LIBRARY_TYPE = 'user' 
# The exact name of the Zotero collection containing your bibliography.
COLLECTION_NAME = 'Name of Your Zotero Collection' 

# --- Document and Font Configuration ---
# The path to your .docx file. Using '~' for the home directory is supported.
# Example: '~/Documents/My_Dissertation.docx'
DOCX_PATH = '~/path/to/your/dissertation.docx'
# The default font to be used for headers and bibliography entries.
FONT_NAME = 'Times New Roman'

# --- Bibliography Category Order and Headers ---
CATEGORY_HEADERS = {
    'note: normative': 'Нормативные и официальные источники',
    'note: dissertation': 'Диссертации и авторефераты',
    'note: academic': 'Научная и учебная литература, статьи из журналов и сборников',
    'note: web': 'Интернет-ресурсы',
    'note: other': 'ПРОЧИЕ ИСТОЧНИКИ'
}
CATEGORY_ORDER = list(CATEGORY_HEADERS.keys())


# ==============================================================================
# --- CORE FUNCTIONS ---
# ==============================================================================

def format_footnotes_xml(docx_path, output_path):
    """
    Directly formats footnotes by manipulating the 'word/footnotes.xml' file
    within the DOCX package.

    This function applies the following formatting to all footnote paragraphs:
    - Font: Times New Roman, 10pt
    - Alignment: Justify
    - Line Spacing: 1.0
    - Indentation: Removes all indents and adds a left tab stop for
      consistent alignment of text after the footnote number.

    Args:
        docx_path (str): The path to the input DOCX file.
        output_path (str): The path to save the modified DOCX file.
    """
    temp_dir = 'temp_docx_footnotes'
    if os.path.exists(temp_dir):
        shutil.rmtree(temp_dir)
    os.makedirs(temp_dir)

    try:
        with zipfile.ZipFile(docx_path, 'r') as zip_ref:
            zip_ref.extractall(temp_dir)

        footnotes_path = os.path.join(temp_dir, 'word', 'footnotes.xml')
        if not os.path.exists(footnotes_path):
            print("Warning: 'footnotes.xml' not found. No footnotes in document.")
            shutil.copy(docx_path, output_path)
            return

        parser = etree.XMLParser(remove_blank_text=True)
        tree = etree.parse(footnotes_path, parser)
        root = tree.getroot()

        ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
        w = '{' + ns['w'] + '}'

        # Apply paragraph-level formatting
        for pPr in root.xpath('.//w:pPr', namespaces=ns):
            # Set Justification
            jc = pPr.find(f'{w}jc')
            if jc is None:
                jc = etree.Element(f'{w}jc', {f'{w}val': 'both'})
                pPr.append(jc)
            else:
                jc.set(f'{w}val', 'both')

            # Set Line Spacing
            spacing = pPr.find(f'{w}spacing')
            if spacing is None:
                spacing = etree.Element(f'{w}spacing', {f'{w}line': '240', f'{w}lineRule': 'auto'})
                pPr.append(spacing)
            else:
                spacing.set(f'{w}line', '240')
                spacing.set(f'{w}lineRule', 'auto')

            # Remove any existing indentation
            ind = pPr.find(f'{w}ind')
            if ind is not None:
                pPr.remove(ind)

            # Add a tab stop for alignment after the footnote number
            tabs = pPr.find(f'{w}tabs')
            if tabs is None:
                tabs = etree.Element(f'{w}tabs')
                pPr.append(tabs)
            tabs.clear()
            tab = etree.Element(f'{w}tab', {f'{w}val': 'left', f'{w}pos': '284'})
            tabs.append(tab)

        # Apply run-level formatting (font and size)
        for rPr in root.xpath('.//w:rPr', namespaces=ns):
            # Set Font
            rFonts = rPr.find(f'{w}rFonts')
            if rFonts is None:
                rFonts = etree.Element(f'{w}rFonts')
                rPr.insert(0, rFonts)
            rFonts.set(f'{w}ascii', FONT_NAME)
            rFonts.set(f'{w}hAnsi', FONT_NAME)
            rFonts.set(f'{w}cs', FONT_NAME)
            rFonts.set(f'{w}eastAsia', FONT_NAME)

            # Set Font Size to 10pt (20 half-points)
            sz = rPr.find(f'{w}sz')
            if sz is None:
                sz = etree.Element(f'{w}sz')
                rPr.append(sz)
            sz.set(f'{w}val', '20')

            szCs = rPr.find(f'{w}szCs')
            if szCs is None:
                szCs = etree.Element(f'{w}szCs')
                rPr.append(szCs)
            szCs.set(f'{w}val', '20')

        tree.write(footnotes_path, xml_declaration=True, encoding='UTF-8', pretty_print=False)

        # Repackage the DOCX file
        with zipfile.ZipFile(output_path, 'w', zipfile.ZIP_DEFLATED) as zip_out:
            for foldername, _, filenames in os.walk(temp_dir):
                for filename in filenames:
                    filepath = os.path.join(foldername, filename)
                    arcname = os.path.relpath(filepath, temp_dir)
                    zip_out.write(filepath, arcname)
        
        print("Footnotes successfully formatted via XML.")

    finally:
        if os.path.exists(temp_dir):
            shutil.rmtree(temp_dir)


def normalize_text(text):
    """
    Cleans and normalizes a string for matching purposes.

    Converts text to lowercase, removes leading/trailing whitespace, and
    replaces non-alphanumeric characters with a single space.

    Args:
        text (str): The input string.

    Returns:
        str: The normalized string.
    """
    text = text.strip().lower()
    text = re.sub(r'[^a-z0-9\sа-яё]', ' ', text)
    text = re.sub(r'\s+', ' ', text).strip()
    return text


def sort_key_bibliography(text):
    """
    Provides a sort key for bibliography entries.

    Sorts entries into two groups: Cyrillic first, then Latin.
    Within each group, sorts alphabetically.

    Args:
        text (str): The bibliography entry text.

    Returns:
        tuple: A tuple (group, sort_string) for sorting.
    """
    cleaned_text = text.strip()
    if not cleaned_text:
        return (3, "")  # Push empty strings to the end

    cyrillic_chars = sum(1 for char in cleaned_text if 'а' <= char.lower() <= 'я' or char.lower() == 'ё')
    latin_chars = sum(1 for char in cleaned_text if 'a' <= char.lower() <= 'z')

    if cyrillic_chars > latin_chars:
        # Group 1: Cyrillic entries, sorted by transliterated value
        return (1, translit(cleaned_text, 'ru', reversed=False).lower())
    else:
        # Group 2: Latin entries, sorted alphabetically
        return (2, cleaned_text.lower())


# ==============================================================================
# --- MAIN SCRIPT LOGIC ---
# ==============================================================================

def main():
    """
    Main function to execute the formatting and sorting process.
    """
    # --- Step 1: Format footnotes ---
    # Expand the user's home directory if '~' is used in the path
    expanded_docx_path = os.path.expanduser(DOCX_PATH)
    if not os.path.exists(expanded_docx_path):
        raise FileNotFoundError(f"The specified file does not exist: {expanded_docx_path}")

    temp_formatted_path = expanded_docx_path.replace('.docx', '_temp.docx')
    format_footnotes_xml(expanded_docx_path, temp_formatted_path)

    # --- Step 2: Prepare Zotero data for bibliography sorting ---
    print("Fetching Zotero data...")
    zot = zotero.Zotero(USER_ID, LIBRARY_TYPE, API_KEY)
    try:
        collections = zot.collections()
    except Exception as e:
        print(f"Error connecting to Zotero: {e}")
        return

    collection_id = next((c['data']['key'] for c in collections if c['data']['name'] == COLLECTION_NAME), None)
    if not collection_id:
        raise RuntimeError(f"Zotero collection '{COLLECTION_NAME}' not found.")

    items = zot.collection_items(collection_id, limit=1000)
    title_category_map = {
        normalize_text(item['data'].get('title', '')): item['data'].get('extra', '').lower()
        for item in items if item['data'].get('title')
    }

    # --- Step 3: Process the document and sort bibliography ---
    print("Processing document and sorting bibliography...")
    doc = Document(temp_formatted_path)

    # Find the start of the bibliography
    start_index = -1
    for i, p in enumerate(doc.paragraphs):
        if "СПИСОК ЛИТЕРАТУРЫ" in p.text.upper():
            start_index = i
            break
    
    if start_index == -1:
        raise RuntimeError("Bibliography heading 'СПИСОК ЛИТЕРАТУРЫ' not found.")

    # Extract and clear old bibliography entries
    entries = [p.text.strip() for p in doc.paragraphs[start_index + 1:] if p.text.strip()]
    for i in range(len(doc.paragraphs) - 1, start_index, -1):
        p = doc.paragraphs[i]
        p._element.getparent().remove(p._element)

    # Classify entries
    grouped_entries = defaultdict(list)
    entry_pattern = re.compile(r"^\d+[\.\)]\s*(.*)")
    for entry in entries:
        text_only = (match.group(1) if (match := entry_pattern.match(entry)) else entry)
        norm_key = normalize_text(text_only)
        
        category_tag = next((cat for t, cat_list in title_category_map.items() if t in norm_key for cat in CATEGORY_ORDER if cat in cat_list), 'note: other')
        grouped_entries[category_tag].append(text_only)

    # Add new, sorted bibliography
    current_bib_index = 1
    first_category_written = False
    for category_tag in CATEGORY_ORDER:
        group = grouped_entries.get(category_tag)
        if not group:
            continue

        sorted_group = sorted(group, key=sort_key_bibliography)

        if first_category_written:
            doc.add_paragraph()  # Add space between categories
        
        header_par = doc.add_paragraph()
        header_par.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = header_par.add_run(CATEGORY_HEADERS[category_tag])
        run.italic = True
        run.font.name = FONT_NAME
        run.font.size = Pt(14)
        header_par.paragraph_format.line_spacing = 1.5
        first_category_written = True

        for item in sorted_group:
            p = doc.add_paragraph(style='List Paragraph')
            p.text = f"{current_bib_index}.\t{item}"
            p_format = p.paragraph_format
            p_format.left_indent = Pt(35.4)  # Approx 1.25 cm
            p_format.first_line_indent = Pt(-35.4)
            p_format.line_spacing = 1.5
            
            # Apply font to the entire paragraph
            for run in p.runs:
                run.font.name = FONT_NAME
                run.font.size = Pt(14)
            
            current_bib_index += 1

    # --- Step 4: Save final document and clean up ---
    output_path = expanded_docx_path.replace('.docx', '_sorted.docx')
    doc.save(output_path)
    os.remove(temp_formatted_path)

    print(f"\nProcessing complete. Final document saved to:\n{output_path}")


if __name__ == '__main__':
    main()
